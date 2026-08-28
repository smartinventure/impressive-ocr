# SPDX-License-Identifier: AGPL-3.0-or-later
"""Page order and page breaks when per-page Word files are merged.

The bug: a 28-page contract came out with 27 blank pages in a row. `add_paragraph()` is a
python-docx API call, and python-docx keeps `sectPr` last in the body and inserts paragraphs
*before* it. The page content was appended with raw lxml, which appends *after* it. So every
page-break paragraph collected together right after page one, and every page's content landed
behind the section properties.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from impressive_ocr_sidecar.writers.page_merge import merge_pages


def page_file(directory: Path, index: int, text: str) -> Path:
    document = Document()
    document.add_paragraph(text)
    path = directory / f"scan_{index}.docx"
    document.save(str(path))
    return path


@pytest.fixture
def pages(tmp_path: Path) -> list[Path]:
    return [page_file(tmp_path, index, f"PAGE{index}") for index in range(4)]


class TestDocxMerge:
    def test_keeps_the_pages_in_order(self, pages: list[Path]) -> None:
        merged = merge_pages(pages, "contract")[0]

        text = [p.text for p in Document(str(merged)).paragraphs if p.text.strip()]

        assert text == ["PAGE0", "PAGE1", "PAGE2", "PAGE3"]

    def test_does_not_bunch_the_page_breaks_together(self, pages: list[Path]) -> None:
        # One empty break paragraph belongs between each pair of pages. All of them in a row
        # is 27 blank pages in a 28-page document, which is what was reported.
        merged = merge_pages(pages, "contract")[0]
        paragraphs = Document(str(merged)).paragraphs

        longest_run = 0
        run = 0
        for paragraph in paragraphs:
            run = run + 1 if not paragraph.text.strip() else 0
            longest_run = max(longest_run, run)

        assert longest_run <= 1

    def test_puts_one_break_between_each_pair_of_pages(self, pages: list[Path]) -> None:
        merged = merge_pages(pages, "contract")[0]

        breaks = sum(
            str(run._element.xml).count('w:type="page"')
            for paragraph in Document(str(merged)).paragraphs
            for run in paragraph.runs
        )

        assert breaks == len(pages) - 1

    def test_leaves_the_section_properties_at_the_end(self, pages: list[Path]) -> None:
        # Content after `sectPr` is out of order for the schema, and the section's page setup
        # then does not govern the pages that follow it.
        merged = merge_pages(pages, "contract")[0]
        body = Document(str(merged)).element.body
        tags = [child.tag.rpartition("}")[2] for child in body]

        assert tags.count("sectPr") == 1
        assert tags[-1] == "sectPr"
