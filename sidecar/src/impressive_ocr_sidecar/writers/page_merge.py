# SPDX-License-Identifier: AGPL-3.0-or-later
"""Join per-page files into one document.

Pages are recognised one at a time — which is what makes progress reporting and cancellation
possible — and PaddleOCR writes a file per page result. A six-page scan therefore produced six
.docx files, which is not what anyone means by "convert this to Word".

Merging happens per format, because "join two of these" means something different each time:
Markdown is text and concatenates, a .docx is a zip of XML whose body elements have to be
moved between documents, and a spreadsheet is neither.

Formats not handled here keep their per-page files, deliberately. A page of JSON is a
structured record of that page and stapling six together produces something no consumer
asked for; per-page .xlsx sheets are the same argument.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from ..core.logging import get_logger

_logger = get_logger()

#: Text formats where concatenation is the whole job.
_TEXT_SUFFIXES = frozenset({".md", ".txt"})

#: Blank line between pages: Markdown needs it to keep the last paragraph of one page from
#: running into the first heading of the next.
_PAGE_SEPARATOR = "\n\n"


def merge_pages(files: list[Path], output_stem: str) -> list[Path]:
    """Combine ``files`` into one document per format where that is meaningful.

    Returns what the caller should report as written: a single merged file, or the originals
    untouched when the format cannot be merged or there is only one of them.
    """
    if len(files) < 2:
        return files

    suffix = files[0].suffix.lower()
    if any(path.suffix.lower() != suffix for path in files):
        # Mixed suffixes mean this is not a set of per-page files (a writer that emits both
        # .html and .htm, say). Leave them alone rather than guessing.
        return files

    if suffix in _TEXT_SUFFIXES:
        return [_merge_text(files, output_stem)]
    if suffix == ".html":
        return [_merge_html(files, output_stem)]
    if suffix == ".docx":
        merged = _merge_docx(files, output_stem)
        return [merged] if merged is not None else files

    return files


def _target(files: list[Path], output_stem: str) -> Path:
    """Where the merged file goes: beside the pages, named for the document."""
    return files[0].parent / f"{output_stem}{files[0].suffix.lower()}"


def _replace_pages(files: list[Path], target: Path) -> Path:
    """Delete the per-page files once their content is safely in ``target``."""
    for path in files:
        if path != target:
            path.unlink(missing_ok=True)
    return target


def _merge_text(files: list[Path], output_stem: str) -> Path:
    target = _target(files, output_stem)
    parts = [path.read_text(encoding="utf-8", errors="replace").strip() for path in files]
    body = _PAGE_SEPARATOR.join(part for part in parts if part)

    # Written to a neighbour first: `target` is usually one of `files`, and truncating it
    # before the later pages have been read would lose them.
    staging = target.with_suffix(target.suffix + ".merging")
    staging.write_text(body + "\n", encoding="utf-8")
    _replace_pages(files, staging)
    staging.replace(target)
    return target


def _merge_html(files: list[Path], output_stem: str) -> Path:
    """Concatenate page bodies inside one document.

    PaddleOCR writes a complete HTML document per page. Concatenating them verbatim yields a
    file with six ``<html>`` elements, which browsers render but no parser should be asked to
    accept, so the bodies are lifted out and wrapped once.
    """
    target = _target(files, output_stem)
    bodies = [_body_of(path.read_text(encoding="utf-8", errors="replace")) for path in files]

    document = (
        "<!doctype html>\n<html>\n<head>\n<meta charset=\"utf-8\">\n"
        f"<title>{output_stem}</title>\n</head>\n<body>\n"
        + "\n<hr>\n".join(bodies)
        + "\n</body>\n</html>\n"
    )

    staging = target.with_suffix(target.suffix + ".merging")
    staging.write_text(document, encoding="utf-8")
    _replace_pages(files, staging)
    staging.replace(target)
    return target


def _body_of(markup: str) -> str:
    """The contents of ``<body>``, or the whole string when there is no body element."""
    lowered = markup.lower()
    start = lowered.find("<body")
    if start == -1:
        return markup.strip()
    start = lowered.find(">", start)
    end = lowered.rfind("</body>")
    if start == -1 or end == -1 or end <= start:
        return markup.strip()
    return markup[start + 1 : end].strip()


def _merge_docx(files: list[Path], output_stem: str) -> Path | None:
    """Append every page's body into the first document, with page breaks between.

    python-docx has no public "merge" operation, so the paragraphs and tables are moved at the
    XML level — which is what its own maintainers recommend and what every recipe for this
    does. The alternative, re-creating each element through the API, loses the table and run
    formatting that made Word output worth producing.
    """
    try:
        import docx
        from docx.enum.text import WD_BREAK
    except ImportError as error:
        _logger.warning(
            "python-docx is unavailable, so the per-page Word files were left separate",
            extra={"error": str(error)},
        )
        return None

    target = _target(files, output_stem)
    staging = target.with_suffix(".merging.docx")
    shutil.copyfile(files[0], staging)

    merged = docx.Document(str(staging))
    body = merged.element.body

    # The section properties are lifted out for the duration of the merge, and that is the
    # whole fix. python-docx keeps `sectPr` last and inserts through the API *before* it,
    # while raw appends land *after* it — so mixing the two put every page-break paragraph in
    # one clump behind page one and every page's content behind the section properties. A
    # 28-page contract came out as its first page, 27 blank pages, then the rest.
    #
    # With it removed both paths append to the end, in the order written, and it goes back
    # last where the schema wants it.
    section_properties = None
    for element in list(body):
        if element.tag.endswith("}sectPr"):
            section_properties = element
            body.remove(element)

    for path in files[1:]:
        merged.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        page = docx.Document(str(path))
        for element in list(page.element.body):
            # A later page's own `sectPr` is dropped rather than copied: it would end the
            # merged document early, at the point that page was appended.
            if element.tag.endswith("}sectPr"):
                continue
            body.append(element)

    if section_properties is not None:
        body.append(section_properties)

    merged.save(str(staging))
    _replace_pages(files, staging)
    staging.replace(target)
    return target
