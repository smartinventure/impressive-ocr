# SPDX-License-Identifier: AGPL-3.0-or-later
"""Per-page files must arrive as one document.

Pages are recognised one at a time, so PaddleOCR writes a file per page: a six-page scan
asked for Word produced six .docx files in the ZIP, which is not what "convert this to Word"
means to anybody.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from impressive_ocr_sidecar.writers.page_merge import merge_pages


def _write(tmp_path: Path, name: str, body: str) -> Path:
    path = tmp_path / name
    path.write_text(body, encoding="utf-8")
    return path


class TestText:
    def test_markdown_pages_become_one_file(self, tmp_path: Path) -> None:
        pages = [
            _write(tmp_path, "scan_0.md", "# Page one\n\nFirst."),
            _write(tmp_path, "scan_1.md", "# Page two\n\nSecond."),
        ]

        merged = merge_pages(pages, "scan")

        assert len(merged) == 1
        assert merged[0].name == "scan.md"
        body = merged[0].read_text(encoding="utf-8")
        assert "First." in body
        assert "Second." in body
        assert body.index("First.") < body.index("Second.")

    def test_the_page_files_are_gone_afterwards(self, tmp_path: Path) -> None:
        # Otherwise the ZIP contains both the merged document and the pages it was made from.
        pages = [
            _write(tmp_path, "scan_0.txt", "one"),
            _write(tmp_path, "scan_1.txt", "two"),
        ]

        merge_pages(pages, "scan")

        assert sorted(p.name for p in tmp_path.iterdir()) == ["scan.txt"]

    def test_a_single_page_is_left_exactly_as_it_was(self, tmp_path: Path) -> None:
        only = _write(tmp_path, "scan_0.md", "# Just the one")

        assert merge_pages([only], "scan") == [only]


class TestHtml:
    def test_bodies_are_lifted_into_one_document(self, tmp_path: Path) -> None:
        pages = [
            _write(tmp_path, "scan_0.html", "<html><body><p>First</p></body></html>"),
            _write(tmp_path, "scan_1.html", "<html><body><p>Second</p></body></html>"),
        ]

        merged = merge_pages(pages, "scan")

        body = merged[0].read_text(encoding="utf-8")
        assert body.lower().count("<html") == 1, "six <html> elements is not a document"
        assert "<p>First</p>" in body
        assert "<p>Second</p>" in body

    def test_a_fragment_without_a_body_element_still_merges(self, tmp_path: Path) -> None:
        pages = [
            _write(tmp_path, "scan_0.html", "<p>First</p>"),
            _write(tmp_path, "scan_1.html", "<p>Second</p>"),
        ]

        body = merge_pages(pages, "scan")[0].read_text(encoding="utf-8")

        assert "<p>First</p>" in body
        assert "<p>Second</p>" in body


class TestDocx:
    def _page(self, tmp_path: Path, name: str, text: str) -> Path:
        path = tmp_path / name
        document = docx.Document()
        document.add_paragraph(text)
        document.save(str(path))
        return path

    def test_six_pages_become_one_word_document(self, tmp_path: Path) -> None:
        pages = [self._page(tmp_path, f"scan_{i}.docx", f"Page {i}") for i in range(6)]

        merged = merge_pages(pages, "scan")

        assert [p.name for p in merged] == ["scan.docx"]
        text = "\n".join(p.text for p in docx.Document(str(merged[0])).paragraphs)
        for index in range(6):
            assert f"Page {index}" in text

    def test_pages_keep_their_order(self, tmp_path: Path) -> None:
        pages = [self._page(tmp_path, f"scan_{i}.docx", f"Page {i}") for i in range(3)]

        merged = merge_pages(pages, "scan")

        text = "\n".join(p.text for p in docx.Document(str(merged[0])).paragraphs)
        assert text.index("Page 0") < text.index("Page 1") < text.index("Page 2")

    def test_tables_survive_the_merge(self, tmp_path: Path) -> None:
        """Re-creating elements through the API would lose exactly this."""
        first = tmp_path / "scan_0.docx"
        document = docx.Document()
        document.add_paragraph("Before")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Net"
        table.cell(0, 1).text = "42.00"
        document.save(str(first))
        second = self._page(tmp_path, "scan_1.docx", "After")

        merged = merge_pages([first, second], "scan")

        result = docx.Document(str(merged[0]))
        assert len(result.tables) == 1
        assert result.tables[0].cell(0, 1).text == "42.00"


class TestFormatsLeftAlone:
    @pytest.mark.parametrize("suffix", [".json", ".xlsx", ".png"])
    def test_unmergeable_formats_keep_their_pages(self, tmp_path: Path, suffix: str) -> None:
        # A page of JSON is a record of that page; stapling six together produces something
        # no consumer asked for.
        pages = [_write(tmp_path, f"scan_{i}{suffix}", "{}") for i in range(2)]

        assert merge_pages(pages, "scan") == pages

    def test_mixed_suffixes_are_not_a_page_set(self, tmp_path: Path) -> None:
        pages = [
            _write(tmp_path, "scan_0.html", "<p>a</p>"),
            _write(tmp_path, "scan_1.htm", "<p>b</p>"),
        ]

        assert merge_pages(pages, "scan") == pages


class TestPageOrdering:
    """Paddle names pages ``scan_0`` … ``scan_10``, which sort as text into 0, 1, 10, 2.

    A merged document whose pages are shuffled is worse than six separate ones, because the
    damage is invisible until somebody reads page eleven.
    """

    def test_a_trailing_page_index_sorts_numerically(self) -> None:
        from impressive_ocr_sidecar.writers.paddle_native import _page_order

        names = [Path(f"scan_{index}.docx") for index in (0, 1, 2, 10, 11)]
        shuffled = [names[3], names[0], names[4], names[2], names[1]]

        assert sorted(shuffled, key=_page_order) == names

    def test_a_name_without_an_index_is_left_to_sort_by_text(self) -> None:
        from impressive_ocr_sidecar.writers.paddle_native import _page_order

        names = [Path("b.docx"), Path("a.docx")]

        assert [p.name for p in sorted(names, key=_page_order)] == ["a.docx", "b.docx"]
