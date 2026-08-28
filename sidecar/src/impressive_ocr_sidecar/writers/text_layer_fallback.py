# SPDX-License-Identifier: AGPL-3.0-or-later
"""Output for pages that were never OCR'd.

A page taken from a PDF's existing text layer has no PaddleOCR result behind it -- there was
no inference, which is the entire point -- so `save_to_markdown` and `save_to_word` have
nothing to act on. The native writer skipped those pages, and a born-digital PDF in `hybrid`
mode therefore finished successfully having written nothing at all: the folder was empty, the
job said "1 succeeded", and the page counter said 0 of 8 because no page needed recognising.

The same class of failure is already recorded in `sidecar/pyproject.toml`, where a missing
`python-docx` meant a job "finished successfully having written everything except the Word
file". Silence is the recurring part, and it is why `paddle_native` now logs a format that
produced nothing rather than returning an empty list.

What this module writes is deliberately plain. The text layer gives us the words and their
boxes and nothing else -- no headings, no table structure -- so inventing a document
hierarchy here would be inventing it from nothing.
"""

from __future__ import annotations

from pathlib import Path

from ..core.logging import get_logger
from ..engines.base import PageResult

_logger = get_logger()

#: Formats that can be produced without a PaddleOCR result.
#:
#: Markdown because the page already carries its own; Word because `python-docx` is a direct
#: dependency and paragraphs are all this content supports anyway. A spreadsheet, an HTML
#: layout or a visualisation image would each need structure the text layer does not have.
FALLBACK_FORMATS = frozenset({"markdown", "docx"})

SUFFIXES = {"markdown": ".md", "docx": ".docx"}


def page_file_name(head: str, page_number: int, suffix: str) -> str:
    """Name a fallback page so it sorts among Paddle's own per-page files.

    `_page_order` in `paddle_native` reads a trailing `_<n>` as a number, so sharing the head
    is what keeps a mixed document -- a digital PDF with scans appended, the case `hybrid`
    exists for -- in reading order instead of grouping the two sources apart.
    """
    return f"{head}_{page_number}{suffix}"


def common_head(paths: list[Path]) -> str | None:
    """The stem Paddle used for its per-page files, without the page index."""
    for path in paths:
        head, separator, tail = path.stem.rpartition("_")
        if separator and tail.isdigit():
            return head
    return None


def write_markdown_page(page: PageResult, target: Path) -> bool:
    """Write one page's Markdown. Returns whether there was anything to write."""
    content = (page.markdown or page.text or "").strip()
    if not content:
        return False
    target.write_text(f"{content}\n", encoding="utf-8")
    return True


def write_docx_page(page: PageResult, target: Path) -> bool:
    """Write one page as a minimal Word document.

    Paragraph per line, with `#` prefixes mapped to headings, because that is the only
    structure the source Markdown actually carries. `python-docx` is a hard dependency, so a
    failure here is a real error rather than a missing optional feature.
    """
    content = (page.markdown or page.text or "").strip()
    if not content:
        return False

    from docx import Document

    document = Document()
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        level = len(stripped) - len(stripped.lstrip("#"))
        if 0 < level <= 6 and stripped[level : level + 1] == " ":
            document.add_heading(stripped[level:].strip(), level=level)
        else:
            document.add_paragraph(stripped)

    document.save(str(target))
    return True


def write_page(page: PageResult, output_format: str, target: Path) -> bool:
    if output_format == "markdown":
        return write_markdown_page(page, target)
    if output_format == "docx":
        return write_docx_page(page, target)
    return False
